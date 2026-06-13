"""
Script sinh file ĐATN_20252_DIEUHOANG.docx
Đồ Án Tốt Nghiệp: Nền Tảng Đọc Truyện Trực Tuyến Tích Hợp AI
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Trang / Margin ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Styles helper ───────────────────────────────────────────────────────────
def set_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East-Asian font as well
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),      "Times New Roman")
    rFonts.set(qn("w:hAnsi"),      "Times New Roman")
    rFonts.set(qn("w:cs"),         "Times New Roman")
    rFonts.set(qn("w:eastAsia"),   "Times New Roman")
    rPr.insert(0, rFonts)

def para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
         italic=False, space_before=0, space_after=6, first_indent=1.0):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    fmt.first_line_indent = Cm(first_indent)
    fmt.line_spacing = Pt(20)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p

def heading1(text, numbering=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    full = f"{numbering} {text}".strip()
    r = p.add_run(full.upper())
    set_font(r, size=14, bold=True)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_font(r, size=13, bold=True)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, italic=True)
    return p

def center_bold(text, size=14, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_font(r, size=size, bold=True)
    return p

def center_text(text, size=13, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p

def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)

def page_break():
    doc.add_page_break()

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1E3A5F")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)
    # data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=11)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ═══════════════════════════════════════════════════════════════════════════
blank(2)
center_text("TRƯỜNG ĐẠI HỌC / HỌC VIỆN", size=13, bold=True)
center_text("KHOA CÔNG NGHỆ THÔNG TIN", size=13, bold=True)
center_text("─────────────────────────────", size=12)
blank(3)
center_bold("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", size=16)
blank(2)
center_bold("NỀN TẢNG ĐỌC TRUYỆN TRỰC TUYẾN\nTÍCH HỢP TRÍ TUỆ NHÂN TẠO", size=18, space_before=0, space_after=4)
blank()
center_text("Online Story Reading Platform with AI Integration", size=14, bold=False)
blank(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
for label, value in [("Sinh viên thực hiện:  ", "Dieu Hoang"),
                     ("Mã sinh viên:         ", ""),
                     ("Lớp:                  ", ""),
                     ("Giáo viên hướng dẫn:  ", ""),
                     ("Khóa:                 ", "20252")]:
    run_l = p.add_run(label)
    set_font(run_l, size=13, bold=True)
    run_v = p.add_run(value + "\n")
    set_font(run_v, size=13)

blank(5)
center_text("Hà Nội, tháng 6 năm 2026", size=13, bold=False)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# LỜI CAM ĐOAN
# ═══════════════════════════════════════════════════════════════════════════
center_bold("LỜI CAM ĐOAN", size=14)
blank()
para(
    "Tôi xin cam đoan đây là công trình nghiên cứu của riêng tôi. "
    "Các số liệu, kết quả trình bày trong báo cáo là trung thực và chưa từng được công bố trong bất kỳ công trình nào khác. "
    "Toàn bộ mã nguồn được xây dựng trong quá trình thực hiện đồ án là do tôi tự thiết kế và lập trình, "
    "có sự hướng dẫn của giáo viên hướng dẫn."
)
para(
    "Tôi xin hoàn toàn chịu trách nhiệm trước Hội đồng bảo vệ về tính trung thực của công trình này."
)
blank(3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("Hà Nội, tháng 6 năm 2026\n\nSinh viên thực hiện\n\n\n\nDieu Hoang")
set_font(r, size=13)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# LỜI CẢM ƠN
# ═══════════════════════════════════════════════════════════════════════════
center_bold("LỜI CẢM ƠN", size=14)
blank()
para(
    "Để hoàn thành đồ án tốt nghiệp này, tôi xin bày tỏ lòng biết ơn sâu sắc đến "
    "các thầy cô trong Khoa Công nghệ Thông tin đã truyền đạt kiến thức nền tảng trong suốt quá trình học tập. "
    "Đặc biệt, tôi xin gửi lời cảm ơn chân thành đến giáo viên hướng dẫn đã tận tình định hướng, "
    "góp ý và hỗ trợ tôi trong từng giai đoạn thực hiện đề tài."
)
para(
    "Tôi cũng gửi lời cảm ơn đến gia đình và bạn bè đã luôn động viên, tạo điều kiện cho tôi hoàn thành tốt đồ án này. "
    "Mặc dù đã cố gắng hết sức, báo cáo chắc chắn còn nhiều thiếu sót. "
    "Tôi rất mong nhận được sự góp ý của thầy cô và hội đồng bảo vệ để hoàn thiện hơn."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TÓM TẮT
# ═══════════════════════════════════════════════════════════════════════════
center_bold("TÓM TẮT", size=14)
blank()
para(
    "Các nền tảng đọc truyện trực tuyến ngày càng phổ biến, nhưng hầu hết vẫn thiếu khả năng tìm kiếm thông minh "
    "và tương tác cá nhân hóa với người dùng. Đồ án này đề xuất và triển khai một nền tảng đọc truyện trực tuyến "
    "tích hợp trí tuệ nhân tạo, giải quyết các hạn chế trên thông qua việc kết hợp nhiều công nghệ hiện đại."
)
para(
    "Hệ thống được xây dựng trên kiến trúc Monorepo với Node.js/Express làm máy chủ, PostgreSQL làm cơ sở dữ liệu quan hệ, "
    "Elasticsearch cho tìm kiếm toàn văn tốc độ cao, và Groq API tích hợp mô hình ngôn ngữ lớn llama-3.1-8b-instant "
    "cho các tính năng AI. Hệ thống crawler tự động thu thập truyện từ các nguồn bên ngoài, "
    "hỗ trợ cả trang tĩnh lẫn trang JavaScript-rendered thông qua cơ chế Axios+Cheerio và Puppeteer. "
    "Tương tác thời gian thực được hiện thực hóa qua Socket.io với cơ chế streaming phản hồi từ AI."
)
para(
    "Đóng góp chính của đồ án gồm: (1) hệ thống chatbot trợ lý nhận thức ngữ cảnh truyện đang đọc, "
    "phản hồi dạng streaming qua Socket.io; (2) cơ chế tìm kiếm ba tầng với khả năng tự động fallback, "
    "đảm bảo hoạt động ổn định ngay cả khi Elasticsearch không khả dụng; "
    "(3) hệ thống gợi ý truyện dựa trên lịch sử đọc cá nhân; "
    "(4) kiến trúc production-ready với phân quyền role-based, bảo mật JWT cookie và parameterized query."
)
para(
    "Kết quả đạt được bao gồm hơn 40 API endpoints hoạt động đầy đủ, 14 bảng cơ sở dữ liệu với chiến lược "
    "indexing tối ưu, và 15 trang giao diện người dùng được xây dựng bằng Vanilla HTML/CSS/JavaScript. "
    "Hệ thống đáp ứng đầy đủ các yêu cầu chức năng và phi chức năng đặt ra, đồng thời sẵn sàng triển khai "
    "trên môi trường production thông qua container hóa bằng Docker."
)

blank()
center_bold("ABSTRACT", size=14)
blank()
para(
    "Online story reading platforms are increasingly popular, yet most lack intelligent search capabilities and "
    "personalized user interaction. This thesis proposes and implements an AI-integrated online story reading platform "
    "that addresses these limitations through a combination of modern web technologies."
)
para(
    "The system is built on a Monorepo architecture with Node.js/Express as the server, PostgreSQL as the relational "
    "database, Elasticsearch for high-performance full-text search, and Groq API integrating the llama-3.1-8b-instant "
    "large language model for AI features. An automated crawler collects stories from external sources, supporting "
    "both static and JavaScript-rendered pages through Axios+Cheerio and Puppeteer mechanisms. "
    "Real-time interaction is implemented via Socket.io with AI response streaming."
)
para(
    "Key contributions include: (1) a context-aware story assistant chatbot that streams responses via Socket.io; "
    "(2) a three-tier search mechanism with automatic fallback ensuring stability even when Elasticsearch is unavailable; "
    "(3) a story recommendation system based on personal reading history; and (4) a production-ready architecture "
    "with role-based access control, JWT cookie security, and parameterized queries."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MỤC LỤC (thủ công)
# ═══════════════════════════════════════════════════════════════════════════
center_bold("MỤC LỤC", size=14)
blank()
toc_items = [
    ("LỜI CAM ĐOAN", ""),
    ("LỜI CẢM ƠN", ""),
    ("TÓM TẮT", ""),
    ("DANH MỤC BẢNG", ""),
    ("DANH MỤC HÌNH ẢNH", ""),
    ("CHƯƠNG 1: GIỚI THIỆU", "1"),
    ("  1.1. Đặt vấn đề", "1"),
    ("  1.2. Mục tiêu đề tài", "2"),
    ("  1.3. Phạm vi nghiên cứu", "3"),
    ("  1.4. Đóng góp của đề tài", "3"),
    ("  1.5. Cấu trúc báo cáo", "4"),
    ("CHƯƠNG 2: KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU", "5"),
    ("  2.1. Khảo sát hiện trạng", "5"),
    ("  2.2. Tổng quan chức năng", "9"),
    ("  2.3. Đặc tả chức năng", "12"),
    ("  2.4. Yêu cầu phi chức năng", "16"),
    ("CHƯƠNG 3: CƠ SỞ LÝ THUYẾT", "17"),
    ("  3.1. Tổng quan về đọc truyện trực tuyến", "17"),
    ("  3.2. Công nghệ backend: Node.js và Express", "18"),
    ("  3.3. Cơ sở dữ liệu quan hệ với PostgreSQL", "19"),
    ("  3.4. Tìm kiếm toàn văn với Elasticsearch", "20"),
    ("  3.5. Mô hình ngôn ngữ lớn và Groq API", "21"),
    ("  3.6. Lập trình thời gian thực với Socket.io", "23"),
    ("  3.7. Bảo mật: JWT và HTTP-only Cookie", "24"),
    ("CHƯƠNG 4: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "25"),
    ("  4.1. Phân tích yêu cầu", "25"),
    ("  4.2. Kiến trúc tổng thể hệ thống", "28"),
    ("  4.3. Thiết kế cơ sở dữ liệu", "31"),
    ("  4.4. Thiết kế API", "37"),
    ("  4.5. Thiết kế giao diện người dùng", "41"),
    ("CHƯƠNG 5: TRIỂN KHAI HỆ THỐNG", "43"),
    ("  5.1. Môi trường phát triển và cấu hình", "43"),
    ("  5.2. Hệ thống thu thập dữ liệu (Crawler)", "44"),
    ("  5.3. Tìm kiếm toàn văn với Elasticsearch", "47"),
    ("  5.4. Tích hợp trí tuệ nhân tạo", "50"),
    ("  5.5. Xác thực và phân quyền người dùng", "55"),
    ("  5.6. Hệ thống thông báo và tương tác thời gian thực", "57"),
    ("  5.7. Triển khai giao diện người dùng", "59"),
    ("CHƯƠNG 6: KIỂM THỬ VÀ ĐÁNH GIÁ", "63"),
    ("  6.1. Môi trường kiểm thử", "63"),
    ("  6.2. Kiểm thử chức năng", "64"),
    ("  6.3. Đánh giá hiệu năng tìm kiếm", "68"),
    ("  6.4. Đánh giá chất lượng chatbot AI", "69"),
    ("  6.5. Nhận xét tổng quan", "70"),
    ("CHƯƠNG 7: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "72"),
    ("  7.1. Kết quả đạt được", "72"),
    ("  7.2. Hạn chế của hệ thống", "73"),
    ("  7.3. Hướng phát triển tiếp theo", "74"),
    ("TÀI LIỆU THAM KHẢO", "75"),
]
for item, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    bold = not item.startswith("  ")
    r = p.add_run(item)
    set_font(r, size=12, bold=bold)
    if pg:
        tab = p.add_run(f"\t{pg}")
        set_font(tab, size=12)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), WD_ALIGN_PARAGRAPH.RIGHT)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# DANH MỤC BẢNG
# ═══════════════════════════════════════════════════════════════════════════
center_bold("DANH MỤC BẢNG", size=14)
blank()
tables_list = [
    ("Bảng 2.1", "So sánh các giải pháp tìm kiếm toàn văn phổ biến"),
    ("Bảng 2.2", "So sánh các nhà cung cấp LLM API"),
    ("Bảng 4.1", "Danh sách yêu cầu chức năng của hệ thống"),
    ("Bảng 4.2", "Danh sách yêu cầu phi chức năng"),
    ("Bảng 4.3", "Mô tả các bảng trong cơ sở dữ liệu"),
    ("Bảng 4.4", "Chiến lược indexing cơ sở dữ liệu"),
    ("Bảng 4.5", "Danh sách API endpoints chính"),
    ("Bảng 5.1", "Biến môi trường cấu hình hệ thống"),
    ("Bảng 5.2", "So sánh chiến lược crawl"),
    ("Bảng 6.1", "Kết quả kiểm thử chức năng xác thực"),
    ("Bảng 6.2", "Kết quả kiểm thử tìm kiếm"),
    ("Bảng 6.3", "Kết quả kiểm thử chatbot AI"),
]
for num, name in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}. ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(name)
    set_font(r2, size=12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# DANH MỤC HÌNH ẢNH
# ═══════════════════════════════════════════════════════════════════════════
center_bold("DANH MỤC HÌNH ẢNH", size=14)
blank()
figures_list = [
    ("Hình 2.1", "Biểu đồ use case tổng quan hệ thống"),
    ("Hình 2.2", "Biểu đồ use case phân rã nhóm Tìm kiếm và Khám phá Truyện"),
    ("Hình 2.3", "Biểu đồ hoạt động quy trình đọc truyện với hỗ trợ AI"),
    ("Hình 5.1", "Kiến trúc tổng thể hệ thống Monorepo"),
    ("Hình 5.2", "Cấu trúc thư mục dự án"),
    ("Hình 5.3", "Sơ đồ quan hệ thực thể (ERD) cơ sở dữ liệu"),
    ("Hình 5.4", "Sơ đồ quan hệ bảng stories – chapters – chapter_contents"),
    ("Hình 5.5", "Luồng xử lý tìm kiếm ba tầng"),
    ("Hình 5.6", "Luồng xử lý chatbot AI qua Socket.io"),
    ("Hình 6.1", "Luồng xử lý crawler với cơ chế fallback Puppeteer"),
    ("Hình 6.2", "Luồng sync dữ liệu PostgreSQL → Elasticsearch"),
    ("Hình 6.3", "Sơ đồ luồng xác thực JWT cookie"),
    ("Hình 6.4", "Giao diện trang chủ hệ thống"),
    ("Hình 6.5", "Giao diện trang đọc truyện với chatbot AI"),
    ("Hình 6.6", "Giao diện trang quản trị Admin"),
    ("Hình 6.1", "Kết quả kiểm thử API với Postman"),
]
for num, name in figures_list:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}. ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(name)
    set_font(r2, size=12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 1: GIỚI THIỆU
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 1: GIỚI THIỆU")

heading2("1.1. Đặt vấn đề")
para(
    "Trong thập kỷ qua, thị trường nội dung số nói chung và đọc truyện trực tuyến nói riêng tăng trưởng "
    "mạnh mẽ tại Việt Nam và trên toàn thế giới. Với sự bùng nổ của smartphone và internet băng thông rộng, "
    "hàng triệu độc giả đã chuyển từ đọc truyện truyền thống sang nền tảng trực tuyến, thúc đẩy nhu cầu "
    "xây dựng các hệ thống phục vụ nội dung quy mô lớn với trải nghiệm người dùng ngày càng cao."
)
para(
    "Tuy nhiên, phần lớn các nền tảng đọc truyện hiện tại vẫn tồn tại những hạn chế đáng kể. "
    "Hệ thống tìm kiếm thường chỉ dừng lại ở đối sánh từ khóa đơn giản, không xử lý được lỗi chính tả "
    "hay tìm kiếm theo ngữ nghĩa. Người dùng mới hoặc người dùng không có định hướng rõ ràng gặp khó khăn "
    "khi muốn khám phá truyện phù hợp với sở thích cá nhân. Hơn nữa, tương tác giữa người đọc và nội dung "
    "còn nghèo nàn — người dùng thiếu công cụ để tóm tắt, đặt câu hỏi hay thảo luận về tác phẩm đang đọc."
)
para(
    "Sự phát triển của các mô hình ngôn ngữ lớn (Large Language Models — LLM) trong những năm gần đây "
    "mở ra cơ hội mới để giải quyết những hạn chế trên. Các LLM có khả năng hiểu ngữ cảnh, tóm tắt văn bản "
    "và trả lời câu hỏi tự nhiên, tạo tiền đề cho việc xây dựng trợ lý ảo đặt trong ngữ cảnh của từng câu chuyện. "
    "Đồng thời, các hệ thống tìm kiếm toàn văn như Elasticsearch cho phép xử lý truy vấn mờ (fuzzy) và "
    "tìm kiếm nhiều trường dữ liệu với độ trễ thấp, vượt xa khả năng của SQL LIKE thông thường."
)
para(
    "Đồ án này xuất phát từ bài toán thực tế: xây dựng một nền tảng đọc truyện trực tuyến hoàn chỉnh, "
    "tích hợp AI và tìm kiếm toàn văn, có khả năng thu thập dữ liệu tự động và sẵn sàng triển khai trên "
    "môi trường production. Hệ thống không chỉ giải quyết nhu cầu cơ bản về lưu trữ và phục vụ nội dung, "
    "mà còn nâng cao trải nghiệm người đọc thông qua các tính năng thông minh được hỗ trợ bởi AI."
)

heading2("1.2. Mục tiêu đề tài")
para(
    "Đề tài đặt ra bốn nhóm mục tiêu chính, mỗi nhóm giải quyết một khía cạnh cụ thể trong chuỗi giá trị "
    "của nền tảng đọc truyện."
)
para(
    "Mục tiêu thứ nhất là xây dựng hệ thống thu thập nội dung tự động. Hệ thống cần có khả năng crawl truyện "
    "từ các nguồn bên ngoài, hỗ trợ cả trang tĩnh lẫn trang sử dụng JavaScript rendering. Quá trình crawl "
    "phải được tối ưu theo cơ chế lazy-load và cache, tránh gọi nguồn bên ngoài lặp lại không cần thiết."
)
para(
    "Mục tiêu thứ hai là triển khai tìm kiếm toàn văn tốc độ cao với khả năng tự phục hồi. "
    "Hệ thống tìm kiếm phải xử lý được lỗi chính tả, hỗ trợ autocomplete, và có cơ chế fallback về "
    "SQL khi Elasticsearch không khả dụng, đảm bảo tính sẵn sàng liên tục của dịch vụ."
)
para(
    "Mục tiêu thứ ba là tích hợp trí tuệ nhân tạo vào trải nghiệm đọc. Cụ thể, hệ thống cần cung cấp "
    "tính năng tóm tắt truyện tự động bằng LLM (có cache để tiết kiệm chi phí API), "
    "và một chatbot trợ lý thời gian thực nhận thức được ngữ cảnh của từng câu chuyện cụ thể mà người "
    "dùng đang đọc."
)
para(
    "Mục tiêu thứ tư là đảm bảo tính production-ready của toàn bộ hệ thống. Kiến trúc phải stateless, "
    "không hardcode thông tin nhạy cảm, có phân quyền role-based đầy đủ, và sẵn sàng container hóa bằng Docker."
)

heading2("1.3. Phạm vi nghiên cứu")
para(
    "Đề tài tập trung vào việc xây dựng một hệ thống web hoàn chỉnh theo mô hình Monorepo, "
    "bao gồm cả backend API và frontend tĩnh trong cùng một kho mã nguồn. "
    "Phạm vi kỹ thuật bao gồm: lớp backend với Node.js và Express v5; lớp lưu trữ với PostgreSQL; "
    "lớp tìm kiếm với Elasticsearch; lớp AI với Groq API; và lớp giao tiếp thời gian thực với Socket.io."
)
para(
    "Đề tài không đặt mục tiêu xây dựng hệ thống phân tán quy mô lớn hay nghiên cứu thuật toán AI từ đầu. "
    "Thay vào đó, trọng tâm là tích hợp hiệu quả các dịch vụ và thư viện hiện có vào một sản phẩm hoàn chỉnh, "
    "đáp ứng yêu cầu của một nền tảng nội dung số quy mô vừa với đầy đủ tính năng."
)

heading2("1.4. Đóng góp của đề tài")
para(
    "Đề tài có ba đóng góp kỹ thuật chính so với các hệ thống tương tự trên thị trường."
)
para(
    "Đóng góp thứ nhất là kiến trúc chatbot nhận thức ngữ cảnh truyện. Thay vì xây dựng một chatbot đa năng chung, "
    "hệ thống tự động nhúng thông tin của truyện đang đọc (tiêu đề, tác giả, thể loại, mô tả, tóm tắt AI) "
    "vào system prompt mỗi phiên hội thoại. Kết hợp với cơ chế lưu lịch sử 20 tin nhắn gần nhất làm context, "
    "chatbot có khả năng trả lời chính xác các câu hỏi liên quan đến nội dung câu chuyện cụ thể."
)
para(
    "Đóng góp thứ hai là cơ chế tìm kiếm ba tầng với tự động fallback. Hệ thống ưu tiên "
    "match_phrase để tìm cụm từ chính xác, chuyển sang multi_match với fuzzy matching khi không tìm thấy, "
    "và cuối cùng fallback về SQL ILIKE nếu Elasticsearch không khả dụng. Thiết kế này đảm bảo "
    "tính sẵn sàng của dịch vụ tìm kiếm trong mọi điều kiện vận hành."
)
para(
    "Đóng góp thứ ba là hệ thống crawl thích ứng. Crawler sử dụng Axios+Cheerio cho trang tĩnh và "
    "tự động chuyển sang Puppeteer cho trang JavaScript-rendered, kết hợp với cơ chế lazy crawl "
    "và cache chương — chỉ thu thập nội dung khi có người đọc lần đầu và lưu vào cơ sở dữ liệu để "
    "phục vụ các lần sau mà không cần crawl lại."
)

heading2("1.5. Cấu trúc báo cáo")
para(
    "Phần còn lại của báo cáo đồ án tốt nghiệp này được tổ chức thành sáu chương, "
    "mỗi chương tập trung vào một khía cạnh cụ thể của quá trình nghiên cứu và phát triển hệ thống."
)
para(
    "Chương 2 thực hiện khảo sát hiện trạng và phân tích yêu cầu — giai đoạn đặt nền móng cho "
    "toàn bộ quá trình thiết kế và triển khai. Chương so sánh các nền tảng đọc truyện hiện có "
    "để xác định điểm hạn chế cần giải quyết, sau đó đặc tả 21 use case dưới dạng biểu đồ use case "
    "và đặc tả chi tiết năm use case quan trọng nhất, cùng các yêu cầu phi chức năng về hiệu năng, "
    "bảo mật và khả năng mở rộng."
)
para(
    "Chương 3 trình bày cơ sở lý thuyết, cung cấp nền tảng học thuật và kỹ thuật cho toàn bộ hệ thống. "
    "Chương lần lượt trình bày từng công nghệ cốt lõi: Node.js và Express với mô hình xử lý sự kiện "
    "bất đồng bộ, PostgreSQL với kiểu dữ liệu JSONB và GIN index, Elasticsearch với inverted index "
    "cho tìm kiếm toàn văn, mô hình ngôn ngữ lớn (LLM) và Groq API, Socket.io cho giao tiếp thời "
    "gian thực, cùng cơ chế bảo mật JWT lưu trong HTTP-only cookie."
)
para(
    "Chương 4 tập trung vào thiết kế kiến trúc và cơ sở dữ liệu hệ thống. Kiến trúc Monorepo "
    "được trình bày với sự phân tách rõ ràng giữa các lớp. Thiết kế cơ sở dữ liệu mô tả 14 bảng "
    "quan hệ cùng hệ thống chỉ mục tối ưu. Chương kết thúc bằng thiết kế tập API REST với 40 "
    "endpoint và phác thảo giao diện người dùng cho hai nhóm trang công khai và riêng tư."
)
para(
    "Chương 5 trình bày đóng góp trọng tâm của đồ án, mô tả chi tiết quá trình hiện thực hóa từng "
    "module. Hệ thống crawler hai tầng (Axios/Cheerio và Puppeteer), tìm kiếm ba tầng với SQL fallback, "
    "chatbot AI streaming qua Socket.io, hệ thống xác thực JWT và thông báo thời gian thực đều "
    "được trình bày chi tiết về thiết kế và cài đặt."
)
para(
    "Chương 6 trình bày quá trình kiểm thử và đánh giá hệ thống. Kiểm thử chức năng theo từng nhóm "
    "use case, đánh giá hiệu năng tìm kiếm và chất lượng chatbot AI là ba nội dung chính. "
    "Phần cuối chương đưa ra nhận xét tổng quan về mức độ đáp ứng các mục tiêu đề ra."
)
para(
    "Chương 7 tổng kết toàn bộ quá trình nghiên cứu và phát triển, đối chiếu kết quả với mục tiêu ban đầu, "
    "phân tích các hạn chế còn tồn tại và đề xuất hướng phát triển tiếp theo."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 2: KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 2: KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU")

para(
    "Chương 1 đã trình bày bối cảnh, mục tiêu và đóng góp của đề tài. "
    "Chương này tập trung vào giai đoạn khảo sát và phân tích yêu cầu — giai đoạn đặt nền móng "
    "cho thiết kế và triển khai hệ thống ở các chương sau. Chương được tổ chức thành bốn phần: "
    "khảo sát hiện trạng và phân tích các hệ thống tương tự trong phần 2.1, "
    "tổng quan chức năng qua biểu đồ use case trong phần 2.2, "
    "đặc tả chi tiết các use case quan trọng trong phần 2.3, "
    "và yêu cầu phi chức năng trong phần 2.4."
)

heading2("2.1. Khảo sát hiện trạng")
para(
    "Khảo sát được tiến hành theo ba nguồn chính: phân tích các hệ thống tương tự đang hoạt động, "
    "quan sát hành vi người dùng trên các nền tảng hiện có, và nghiên cứu xu hướng công nghệ "
    "trong lĩnh vực phân phối nội dung số. Phần này tập trung vào nguồn thứ nhất — "
    "phân tích và so sánh các nền tảng đọc truyện trực tuyến tiêu biểu tại Việt Nam và quốc tế."
)
para(
    "Thị trường đọc truyện trực tuyến tại Việt Nam hiện có hai nhóm nền tảng chính. "
    "Nhóm thứ nhất là các trang web thuần Việt như TruyenFull.vn và Truyencv.com, "
    "hoạt động chủ yếu dựa trên việc tổng hợp nội dung từ cộng đồng và tự động thu thập từ nguồn ngoài. "
    "Nhóm thứ hai là các nền tảng quốc tế như Wattpad và MangaDex, "
    "hỗ trợ đa ngôn ngữ và có cơ sở hạ tầng phát triển hơn nhưng không tối ưu cho nội dung tiếng Việt."
)
para(
    "Phân tích so sánh cho thấy cả bốn nền tảng đều thiếu khả năng tích hợp AI cho việc hỗ trợ độc giả. "
    "TruyenFull.vn và Truyencv.com sử dụng tìm kiếm theo từ khóa đơn giản (SQL LIKE), "
    "không xử lý được lỗi chính tả hay tìm kiếm đa trường. Wattpad có hệ thống gợi ý nhưng "
    "dựa trên hành vi đọc tập thể, không cá nhân hóa theo lịch sử của từng người dùng. "
    "MangaDex tuy cung cấp API mạnh mẽ nhưng không có tính năng AI hay chatbot. "
    "Bảng 2.1 tóm tắt kết quả so sánh chi tiết."
)
add_table(
    ["Tính năng", "TruyenFull", "Truyencv", "Wattpad", "MangaDex", "DH.Story"],
    [
        ["Tìm kiếm toàn văn (fuzzy)", "Không", "Không", "Có", "Cơ bản", "Có (3 tầng)"],
        ["Autocomplete gợi ý", "Không", "Không", "Có", "Không", "Có"],
        ["Tóm tắt nội dung AI", "Không", "Không", "Không", "Không", "Có"],
        ["Chatbot trợ lý AI", "Không", "Không", "Không", "Không", "Có (streaming)"],
        ["Gợi ý cá nhân hóa", "Cơ bản", "Cơ bản", "Có", "Không", "Có"],
        ["Đọc truyện tranh (ảnh)", "Có", "Có", "Không", "Có", "Có"],
        ["Thông báo thời gian thực", "Không", "Không", "Có", "Không", "Có"],
        ["Bình luận dạng cây", "Có", "Có", "Có", "Không", "Có"],
        ["Báo lỗi và phản hồi admin", "Không", "Không", "Cơ bản", "Có", "Có"],
    ],
    col_widths=[4.5, 2.0, 2.0, 2.0, 2.0, 2.5]
)
para(
    "Từ kết quả khảo sát, ba điểm hạn chế nổi bật của các hệ thống hiện có được xác định làm "
    "định hướng phát triển của đồ án. Thứ nhất, không có nền tảng nào tích hợp AI trực tiếp vào "
    "trải nghiệm đọc — độc giả không thể hỏi về nội dung truyện hay nhận tóm tắt tự động. "
    "Thứ hai, chất lượng tìm kiếm của các nền tảng tiếng Việt còn yếu, đặc biệt khi người dùng "
    "nhập từ khóa không chính xác hoặc không nhớ chính xác tên truyện. "
    "Thứ ba, tính năng gợi ý truyện trên các nền tảng Việt Nam còn đơn giản, chưa dựa trên "
    "phân tích sâu lịch sử đọc của từng người dùng."
)

heading2("2.2. Tổng quan chức năng")
heading3("2.2.1. Biểu đồ use case tổng quan")
para(
    "Hệ thống có ba tác nhân chính tương tác với ba nhóm chức năng khác nhau. "
    "Khách (Guest) là người dùng chưa đăng nhập, có thể xem trang chủ, tìm kiếm truyện, "
    "xem chi tiết và đọc truyện ở chế độ giới hạn. "
    "Người dùng (User) là người đã đăng nhập với tài khoản hợp lệ, được truy cập đầy đủ "
    "tất cả tính năng bao gồm AI, bình luận, đánh giá và quản lý danh sách yêu thích. "
    "Quản trị viên (Admin) kế thừa toàn bộ quyền của người dùng và thêm quyền quản lý "
    "toàn bộ hệ thống: truyện, tài khoản người dùng, báo lỗi và thống kê."
)
para(
    "Biểu đồ use case tổng quan (Hình 2.1) thể hiện sáu nhóm chức năng chính. "
    "Nhóm Quản lý tài khoản bao gồm đăng ký, đăng nhập, quên mật khẩu và cập nhật thông tin cá nhân. "
    "Nhóm Khám phá truyện bao gồm tìm kiếm toàn văn, gợi ý autocomplete, lọc theo thể loại "
    "và xem danh sách truyện mới nhất hoặc phổ biến nhất trong tuần. "
    "Nhóm Đọc truyện cho phép xem nội dung từng chương với cơ chế lazy crawl và cache tự động. "
    "Nhóm Tương tác bao gồm bình luận dạng cây, đánh giá sao, quản lý danh sách yêu thích "
    "và gửi báo lỗi. Nhóm Tính năng AI cung cấp chatbot trợ lý theo ngữ cảnh truyện "
    "và tóm tắt nội dung tự động. Nhóm Quản trị dành riêng cho admin để quản lý hệ thống."
)
para("[Hình 2.1: Biểu đồ use case tổng quan hệ thống]", first_indent=0.0)

heading3("2.2.2. Biểu đồ use case phân rã: Tìm kiếm và Khám phá Truyện")
para(
    "Tìm kiếm và khám phá truyện là nhóm chức năng cốt lõi, phục vụ cả ba nhóm tác nhân "
    "và là điểm tiếp xúc đầu tiên của người dùng với hệ thống. "
    "Nhóm này được phân rã thành năm use case cụ thể: tìm kiếm theo từ khóa, "
    "nhận gợi ý autocomplete khi nhập liệu, lọc truyện theo thể loại, "
    "xem danh sách truyện mới cập nhật, và xem danh sách truyện phổ biến trong tuần."
)
para(
    "Ba use case đầu tiên được kết nối với nhau qua quan hệ include: "
    "khi người dùng thực hiện tìm kiếm, hệ thống đồng thời trả về gợi ý autocomplete "
    "theo thời gian thực. Use case tìm kiếm mở rộng thêm hai trường hợp ngoại lệ: "
    "không tìm thấy kết quả (hiển thị thông báo) và Elasticsearch không khả dụng "
    "(tự động fallback sang SQL). Biểu đồ phân rã được thể hiện tại Hình 2.2."
)
para("[Hình 2.2: Biểu đồ use case phân rã nhóm Tìm kiếm và Khám phá Truyện]", first_indent=0.0)

heading3("2.2.3. Quy trình nghiệp vụ")
para(
    "Quy trình nghiệp vụ tiêu biểu của hệ thống là luồng 'Đọc Truyện với Hỗ Trợ AI', "
    "kết hợp nhiều use case để tạo ra một trải nghiệm đọc đầy đủ. "
    "Quy trình này gồm năm bước chính, được minh họa qua biểu đồ hoạt động tại Hình 2.3."
)
para(
    "Bước một, người dùng tìm kiếm truyện theo từ khóa hoặc lọc theo thể loại. "
    "Bước hai, sau khi chọn truyện, hệ thống hiển thị trang chi tiết và tự động tạo tóm tắt AI — "
    "nếu bản tóm tắt chưa có trong cơ sở dữ liệu thì gọi Groq API và lưu lại, "
    "ngược lại trả về kết quả từ cache. "
    "Bước ba, người dùng chọn chương để đọc; hệ thống kiểm tra cache trong bảng chapter_contents "
    "và crawl ảnh từ nguồn nếu chưa có. "
    "Bước bốn, trong khi đọc, người dùng có thể để lại bình luận hoặc đánh giá sao. "
    "Bước năm, người dùng mở widget chatbot để đặt câu hỏi về nội dung truyện; "
    "hệ thống gọi Groq API với ngữ cảnh truyện và phát trực tuyến câu trả lời qua Socket.io."
)
para("[Hình 2.3: Biểu đồ hoạt động quy trình đọc truyện với hỗ trợ AI]", first_indent=0.0)

heading2("2.3. Đặc tả chức năng")
para(
    "Trong số 21 use case đã xác định, năm use case sau được chọn để đặc tả chi tiết "
    "do mức độ quan trọng và độ phức tạp của luồng xử lý: tìm kiếm truyện (UC08), "
    "đọc chương truyện (UC07), trò chuyện với chatbot AI (UC11), "
    "tóm tắt nội dung truyện bằng AI (UC10), và báo lỗi truyện (UC20)."
)

heading3("2.3.1. Đặc tả use case UC08: Tìm kiếm truyện")
add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Mã use case", "UC08"],
        ["Tên use case", "Tìm kiếm truyện"],
        ["Tác nhân", "Khách, Người dùng, Quản trị viên"],
        ["Tiền điều kiện", "Hệ thống đang hoạt động bình thường"],
        ["Hậu điều kiện", "Danh sách truyện phù hợp được hiển thị trên màn hình"],
        ["Luồng sự kiện chính",
         "1. Người dùng nhập từ khóa vào ô tìm kiếm\n"
         "2. Hệ thống trả về gợi ý autocomplete theo thời gian thực từ Elasticsearch\n"
         "3. Người dùng chọn gợi ý hoặc nhấn Enter để tìm kiếm đầy đủ\n"
         "4. Hệ thống thực hiện tìm kiếm theo ba tầng ưu tiên:\n"
         "   - Tầng 1: match_phrase (Elasticsearch)\n"
         "   - Tầng 2: multi_match fuzzy (Elasticsearch)\n"
         "   - Tầng 3: SQL ILIKE fallback (PostgreSQL)\n"
         "5. Kết quả được trả về theo thứ tự liên quan và hiển thị dạng thẻ truyện"],
        ["Luồng sự kiện phụ",
         "4a. Elasticsearch không khả dụng: hệ thống tự động chuyển sang SQL ILIKE\n"
         "5a. Không có kết quả: hiển thị thông báo 'Không tìm thấy truyện phù hợp'"],
    ],
    col_widths=[4.0, 10.0]
)

heading3("2.3.2. Đặc tả use case UC07: Đọc chương truyện")
add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Mã use case", "UC07"],
        ["Tên use case", "Đọc chương truyện"],
        ["Tác nhân", "Khách, Người dùng"],
        ["Tiền điều kiện", "Truyện và chương tương ứng đã tồn tại trong hệ thống"],
        ["Hậu điều kiện",
         "Nội dung chương được hiển thị;\n"
         "nếu người dùng đã đăng nhập thì lịch sử xem được ghi nhận"],
        ["Luồng sự kiện chính",
         "1. Người dùng chọn chương từ danh sách chương\n"
         "2. Hệ thống gọi GET /api/chapters/:id/content\n"
         "3. Hệ thống kiểm tra bảng chapter_contents:\n"
         "   - Nếu đã có cache: trả về danh sách URL ảnh ngay lập tức\n"
         "   - Nếu chưa có: crawl ảnh từ source_url, lưu vào cache, trả về\n"
         "4. Client hiển thị tuần tự các trang ảnh của chương\n"
         "5. Ghi nhận lịch sử xem (nếu người dùng đã đăng nhập)"],
        ["Luồng sự kiện phụ",
         "3a. Crawl thất bại (network error): trả về HTTP 500 với thông báo lỗi chung\n"
         "3b. Chương không có ảnh: trả về mảng rỗng, client hiển thị thông báo tương ứng"],
    ],
    col_widths=[4.0, 10.0]
)

heading3("2.3.3. Đặc tả use case UC11: Trò chuyện với chatbot AI")
add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Mã use case", "UC11"],
        ["Tên use case", "Trò chuyện với chatbot AI"],
        ["Tác nhân", "Người dùng (bắt buộc đăng nhập)"],
        ["Tiền điều kiện",
         "Người dùng đã đăng nhập và đang ở trang đọc truyện (read2.html hoặc chapter.html)"],
        ["Hậu điều kiện",
         "Câu trả lời của AI hiển thị theo luồng streaming;\n"
         "cả câu hỏi và câu trả lời được lưu vào bảng chat_messages"],
        ["Luồng sự kiện chính",
         "1. Người dùng nhấn mở widget chatbot góc dưới phải\n"
         "2. Hệ thống tải 20 tin nhắn gần nhất qua GET /api/chat/history\n"
         "3. Người dùng nhập câu hỏi và gửi\n"
         "4. Client emit sự kiện chatMessage qua Socket.io kèm nội dung và story_id\n"
         "5. Server kiểm tra cooldown 2 giây; nếu OK thì gọi Groq API với:\n"
         "   - System prompt chứa ngữ cảnh truyện (tiêu đề, tác giả, thể loại, tóm tắt AI)\n"
         "   - Lịch sử 20 tin nhắn gần nhất\n"
         "6. Server emit từng chatChunk khi nhận được từ Groq API (streaming)\n"
         "7. Server emit chatDone khi hoàn thành; lưu tin nhắn vào chat_messages"],
        ["Luồng sự kiện phụ",
         "5a. Cooldown chưa hết (< 2s): server emit chatError 'Vui lòng chờ'\n"
         "5b. Groq API lỗi: server emit chatError với thông báo lỗi chung"],
    ],
    col_widths=[4.0, 10.0]
)

heading3("2.3.4. Đặc tả use case UC10: Tóm tắt nội dung truyện bằng AI")
add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Mã use case", "UC10"],
        ["Tên use case", "Tóm tắt nội dung truyện bằng AI"],
        ["Tác nhân", "Người dùng, Quản trị viên"],
        ["Tiền điều kiện", "Truyện tồn tại trong hệ thống với ít nhất tiêu đề và mô tả"],
        ["Hậu điều kiện",
         "Nội dung tóm tắt hiển thị trên trang chi tiết truyện;\n"
         "kết quả lưu vào trường ai_summary trong bảng stories để tái sử dụng"],
        ["Luồng sự kiện chính",
         "1. Người dùng xem trang chi tiết truyện hoặc nhấn nút 'Tóm tắt bằng AI'\n"
         "2. Client gọi POST /api/ai/summarize với story_id\n"
         "3. Hệ thống kiểm tra trường ai_summary trong bảng stories:\n"
         "   - Nếu đã có: trả về nội dung tóm tắt đã lưu (không gọi AI)\n"
         "   - Nếu chưa có: tiếp tục bước 4\n"
         "4. Hệ thống gọi Groq API (model llama-3.1-8b-instant, temperature 0.3)\n"
         "   với prompt chứa title, genres, description của truyện\n"
         "5. Lưu kết quả vào cột ai_summary và trả về cho client"],
        ["Luồng sự kiện phụ",
         "4a. Groq API lỗi hoặc timeout: trả về HTTP 500 với thông báo lỗi chung"],
    ],
    col_widths=[4.0, 10.0]
)

heading3("2.3.5. Đặc tả use case UC20: Báo lỗi truyện")
add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Mã use case", "UC20"],
        ["Tên use case", "Báo lỗi truyện"],
        ["Tác nhân", "Người dùng"],
        ["Tiền điều kiện", "Người dùng đang xem trang truyện hoặc chương truyện có vấn đề"],
        ["Hậu điều kiện",
         "Báo lỗi lưu vào bảng reports với trạng thái pending;\n"
         "quản trị viên được thông báo qua Socket.io nếu đang online"],
        ["Luồng sự kiện chính",
         "1. Người dùng nhấn nút 'Báo lỗi' trên trang truyện\n"
         "2. Form báo lỗi hiển thị với các trường: tiêu đề, URL trang lỗi,\n"
         "   mô tả vấn đề (bắt buộc), ảnh chụp màn hình (tùy chọn)\n"
         "3. Người dùng điền thông tin và nhấn Gửi\n"
         "4. Client gọi POST /api/report\n"
         "5. Hệ thống lưu báo lỗi với status='pending'\n"
         "6. Client hiển thị xác nhận 'Báo lỗi đã được ghi nhận'"],
        ["Luồng sự kiện phụ",
         "3a. Không điền mô tả: hiển thị cảnh báo validation\n"
         "3b. Ảnh quá lớn (> 5MB): hiển thị cảnh báo và yêu cầu chọn lại"],
    ],
    col_widths=[4.0, 10.0]
)

heading2("2.4. Yêu cầu phi chức năng")
para(
    "Hệ thống cần đáp ứng các yêu cầu phi chức năng thuộc bốn nhóm chính: "
    "hiệu năng, bảo mật, tính khả dụng và khả năng mở rộng."
)
para(
    "Về hiệu năng, API tìm kiếm với Elasticsearch cần phản hồi trong vòng 200ms; "
    "khi fallback sang SQL thì không quá 1 giây. Trang chương truyện lần đầu truy cập "
    "(cần crawl) cho phép tối đa 5 giây; lần sau từ cache không quá 500ms. "
    "Chatbot cần emit chunk đầu tiên trong vòng 2 giây từ khi nhận yêu cầu."
)
para(
    "Về bảo mật, JWT được lưu trong HTTP-only cookie để phòng chống XSS. "
    "Mọi truy vấn database đều sử dụng parameterized query, không nối chuỗi SQL "
    "để phòng chống SQL injection. Thông tin lỗi chi tiết không được lộ ra client — "
    "chỉ trả về message chung và ghi chi tiết vào server log. "
    "Mọi API admin đều có middleware xác thực JWT và kiểm tra role."
)
para(
    "Về tính khả dụng, hệ thống tìm kiếm phải hoạt động ngay cả khi Elasticsearch không khả dụng "
    "nhờ cơ chế SQL fallback. Về khả năng mở rộng, server không lưu trạng thái trong bộ nhớ "
    "giữa các request (stateless), cho phép chạy nhiều instance đồng thời sau load balancer. "
    "Toàn bộ cấu hình được đọc từ biến môi trường, tách biệt hoàn toàn giữa môi trường development "
    "và production."
)
para(
    "Chương này đã phân tích hiện trạng các hệ thống tương tự, xác định ba điểm hạn chế chính "
    "cần giải quyết, và đặc tả 21 use case cùng bốn nhóm yêu cầu phi chức năng. "
    "Trên nền tảng yêu cầu đã phân tích, Chương 3 sẽ trình bày các nền tảng lý thuyết và công nghệ "
    "được lựa chọn để đáp ứng những yêu cầu này."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3: CƠ SỞ LÝ THUYẾT
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 3: CƠ SỞ LÝ THUYẾT")

heading2("3.1. Tổng quan về đọc truyện trực tuyến")
para(
    "Các nền tảng đọc truyện trực tuyến đang hội tụ ba xu hướng công nghệ: nội dung do người dùng sáng tạo (UGC), "
    "thu thập nội dung tự động (web crawling), và cá nhân hóa trải nghiệm thông qua AI. "
    "Việc kết hợp ba xu hướng này đặt ra thách thức kỹ thuật đáng kể về hiệu năng, tính nhất quán dữ liệu "
    "và khả năng mở rộng."
)
para(
    "Về mặt kiến trúc, các hệ thống nội dung quy mô lớn thường áp dụng mô hình tách biệt giữa "
    "lớp thu thập (ingestion), lưu trữ (storage), tìm kiếm (search) và phục vụ (serving). "
    "Mỗi lớp có yêu cầu kỹ thuật riêng: lớp thu thập cần xử lý đa dạng định dạng nguồn; "
    "lớp lưu trữ cần đảm bảo toàn vẹn quan hệ; lớp tìm kiếm cần độ trễ thấp và độ chính xác cao; "
    "lớp phục vụ cần hỗ trợ đồng thời nhiều người dùng."
)

heading2("3.2. Công nghệ backend: Node.js và Express")
para(
    "Node.js là môi trường chạy JavaScript phía máy chủ, xây dựng trên engine V8 của Google Chrome. "
    "Kiến trúc event-driven, non-blocking I/O của Node.js phù hợp đặc biệt cho các ứng dụng yêu cầu "
    "xử lý nhiều kết nối đồng thời với độ trễ thấp — đặc điểm quan trọng khi phục vụ nhiều người dùng đọc truyện "
    "cùng lúc và duy trì kết nối WebSocket cho chatbot."
)
para(
    "Express là framework web tối giản cho Node.js, cung cấp cơ chế routing, middleware và xử lý request/response "
    "linh hoạt. Phiên bản Express v5 được sử dụng trong đề tài cải thiện xử lý lỗi bất đồng bộ "
    "so với v4 — async/await trong route handler có thể ném lỗi mà không cần try/catch bọc ngoài, "
    "Express v5 tự động chuyển lỗi đến error handler."
)
para(
    "Mô hình middleware của Express cho phép xử lý cross-cutting concerns như xác thực JWT, "
    "logging và CORS một cách tập trung. Middleware được thực thi tuần tự theo thứ tự đăng ký, "
    "tạo ra pipeline xử lý request rõ ràng và dễ kiểm soát."
)

heading2("3.3. Cơ sở dữ liệu quan hệ với PostgreSQL")
para(
    "PostgreSQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở với tính năng phong phú và tuân thủ "
    "chuẩn SQL nghiêm ngặt. PostgreSQL hỗ trợ kiểu dữ liệu JSONB cho phép lưu trữ dữ liệu bán cấu trúc "
    "với khả năng index và query hiệu quả — tính năng quan trọng khi lưu trữ danh sách URL ảnh chương truyện tranh."
)
para(
    "Cơ chế index của PostgreSQL đa dạng, bao gồm B-tree cho các truy vấn phạm vi và sắp xếp, "
    "GIN (Generalized Inverted Index) cho array và full-text search, và partial index để tối ưu "
    "cho các điều kiện lọc phổ biến. Đề tài tận dụng GIN index trên cột genres kiểu text[] "
    "để tìm kiếm truyện theo thể loại với độ phức tạp O(log n)."
)
para(
    "Module pg (node-postgres) cung cấp kết nối đến PostgreSQL từ Node.js thông qua cơ chế connection pool. "
    "Pool quản lý một tập kết nối tái sử dụng, giảm overhead khởi tạo kết nối TCP cho mỗi query. "
    "Parameterized query trong pg ngăn chặn SQL injection bằng cách tách biệt hoàn toàn code SQL "
    "và dữ liệu người dùng."
)

heading2("3.4. Tìm kiếm toàn văn với Elasticsearch")
para(
    "Elasticsearch là công cụ tìm kiếm và phân tích phân tán, xây dựng trên nền Apache Lucene. "
    "Elasticsearch index dữ liệu dưới dạng inverted index — cấu trúc dữ liệu ánh xạ từng term "
    "đến danh sách các document chứa term đó — cho phép tìm kiếm toàn văn với độ trễ dưới một giây "
    "trên tập dữ liệu hàng triệu documents."
)
para(
    "Bảng 2.1 so sánh Elasticsearch với các giải pháp tìm kiếm khác để thấy rõ lý do lựa chọn."
)
add_table(
    ["Tiêu chí", "Elasticsearch", "PostgreSQL FTS", "Meilisearch"],
    [
        ["Fuzzy search", "Tốt (Levenshtein)", "Hạn chế", "Tốt"],
        ["Autocomplete", "Tốt (prefix/ngram)", "Hạn chế", "Tốt"],
        ["Scalability", "Rất tốt (phân tán)", "Trung bình", "Trung bình"],
        ["Vietnamese analyzer", "Cần cấu hình thêm", "Kém", "Trung bình"],
        ["Fallback khi down", "Cần xử lý thủ công", "Không cần", "Cần xử lý"],
    ],
    col_widths=[4, 4, 4, 4]
)
para(
    "Elasticsearch hỗ trợ nhiều loại truy vấn linh hoạt. match_phrase tìm kiếm cụm từ chính xác "
    "theo thứ tự xuất hiện; multi_match tìm kiếm đồng thời trên nhiều trường; "
    "fuzzy matching cho phép sai số trong từ khóa dựa trên khoảng cách Levenshtein; "
    "match_phrase_prefix hỗ trợ autocomplete bằng cách khớp prefix của từ cuối cùng trong cụm từ."
)

heading2("3.5. Mô hình ngôn ngữ lớn và Groq API")
para(
    "Mô hình ngôn ngữ lớn (Large Language Model — LLM) là các mạng neural nhân tạo được huấn luyện "
    "trên lượng văn bản khổng lồ, có khả năng sinh văn bản, trả lời câu hỏi và thực hiện các tác vụ "
    "xử lý ngôn ngữ tự nhiên phức tạp. Kiến trúc Transformer với cơ chế self-attention cho phép LLM "
    "nắm bắt quan hệ ngữ nghĩa trong văn bản dài."
)
para(
    "Groq là nhà cung cấp dịch vụ inference LLM tốc độ cao, sử dụng chip LPU (Language Processing Unit) "
    "chuyên dụng. Groq API cung cấp giao diện tương thích với OpenAI API, cho phép sử dụng các mô hình "
    "mã nguồn mở như Llama của Meta với latency thấp hơn đáng kể so với các nhà cung cấp GPU-based thông thường."
)
para(
    "Bảng 2.2 so sánh các nhà cung cấp LLM API chính để làm rõ lý do chọn Groq."
)
add_table(
    ["Nhà cung cấp", "Mô hình", "Tốc độ (token/s)", "Chi phí ($/1M token)", "Streaming"],
    [
        ["Groq", "llama-3.1-8b-instant", "~800+", "Miễn phí (trial)", "Có"],
        ["OpenAI", "gpt-4o-mini", "~100-200", "~0.15 input", "Có"],
        ["Anthropic", "claude-haiku", "~100-150", "~0.25 input", "Có"],
        ["Together AI", "llama-3.1-8b", "~100-200", "~0.18 input", "Có"],
    ],
    col_widths=[3.5, 4, 3, 4, 2]
)
para(
    "Mô hình llama-3.1-8b-instant của Meta, được phục vụ qua Groq, cung cấp chất lượng phản hồi "
    "đủ tốt cho tác vụ tóm tắt và trả lời câu hỏi bối cảnh đơn, đồng thời có tốc độ inference "
    "cao nhất trong số các lựa chọn — yếu tố quan trọng để đảm bảo trải nghiệm streaming mượt mà."
)
para(
    "Kỹ thuật prompt engineering được áp dụng để kiểm soát chất lượng đầu ra. System prompt cung cấp "
    "cho mô hình ngữ cảnh về vai trò (trợ lý đọc truyện) và thông tin truyện cụ thể, "
    "trong khi temperature thấp (0.3 cho tóm tắt) đảm bảo đầu ra nhất quán và ít sáng tạo tùy tiện."
)

heading2("3.6. Lập trình thời gian thực với Socket.io")
para(
    "WebSocket là giao thức truyền thông hai chiều, full-duplex trên nền TCP, cho phép máy chủ "
    "chủ động đẩy dữ liệu đến client mà không cần client gửi request trước. "
    "So với polling HTTP truyền thống, WebSocket giảm đáng kể độ trễ và overhead băng thông, "
    "phù hợp cho các ứng dụng yêu cầu phản hồi tức thời như chat và thông báo."
)
para(
    "Socket.io là thư viện xây dựng trên WebSocket, bổ sung thêm các tính năng như tự động fallback "
    "về long-polling khi WebSocket không khả dụng, room-based broadcasting (phát tin đến nhóm client cụ thể), "
    "và event-driven API thống nhất giữa server và client. Khái niệm 'room' trong Socket.io "
    "được tận dụng để định tuyến thông báo đến đúng người dùng theo địa chỉ email."
)

heading2("3.7. Bảo mật: JWT và HTTP-only Cookie")
para(
    "JSON Web Token (JWT) là tiêu chuẩn mã hóa thông tin giữa các bên dưới dạng đối tượng JSON được ký số. "
    "JWT gồm ba phần: header (thuật toán ký), payload (claims — dữ liệu người dùng) và signature (chữ ký xác thực). "
    "Máy chủ ký JWT bằng secret key; client gửi JWT kèm mỗi request; máy chủ xác minh chữ ký "
    "để tin tưởng thông tin trong payload mà không cần truy vấn database."
)
para(
    "HTTP-only cookie là cookie không thể truy cập từ JavaScript phía client (document.cookie). "
    "Lưu JWT trong HTTP-only cookie — thay vì localStorage — ngăn chặn tấn công Cross-Site Scripting (XSS) "
    "đánh cắp token: ngay cả khi kẻ tấn công inject được JavaScript vào trang, "
    "script đó cũng không thể đọc được cookie chứa JWT. "
    "Kết hợp với thuộc tính SameSite, cơ chế này cũng giảm thiểu nguy cơ CSRF."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 4: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

heading2("4.1. Phân tích yêu cầu")
heading3("4.1.1. Yêu cầu chức năng")
para(
    "Hệ thống cần đáp ứng các yêu cầu chức năng thuộc bảy nhóm chính. "
    "Bảng 4.1 liệt kê đầy đủ các yêu cầu cùng mức độ ưu tiên."
)
add_table(
    ["Mã", "Yêu cầu chức năng", "Ưu tiên"],
    [
        ["UC01", "Đăng ký, đăng nhập, đăng xuất tài khoản", "Cao"],
        ["UC02", "Quên mật khẩu qua OTP email", "Cao"],
        ["UC03", "Cập nhật thông tin cá nhân và upload avatar", "Trung bình"],
        ["UC04", "Phân quyền role-based (user/admin)", "Cao"],
        ["UC05", "Thu thập truyện tự động từ nguồn ngoài (crawler)", "Cao"],
        ["UC06", "CRUD truyện và chương", "Cao"],
        ["UC07", "Đọc truyện theo chương với lazy crawl và cache", "Cao"],
        ["UC08", "Tìm kiếm toàn văn với autocomplete và fallback", "Cao"],
        ["UC09", "Lọc truyện theo thể loại", "Trung bình"],
        ["UC10", "Tóm tắt truyện bằng AI (có cache)", "Cao"],
        ["UC11", "Chatbot trợ lý thời gian thực nhận thức ngữ cảnh", "Cao"],
        ["UC12", "Gợi ý truyện dựa trên lịch sử đọc", "Trung bình"],
        ["UC13", "Bình luận dạng cây và reply", "Trung bình"],
        ["UC14", "Like/unlike bình luận", "Thấp"],
        ["UC15", "Đánh giá truyện (1–5 sao)", "Trung bình"],
        ["UC16", "Quản lý danh sách truyện yêu thích", "Trung bình"],
        ["UC17", "Ghi lịch sử xem truyện", "Cao"],
        ["UC18", "Thống kê tổng quan và truyện hot tuần (Admin)", "Cao"],
        ["UC19", "Quản lý tài khoản người dùng (Admin)", "Cao"],
        ["UC20", "Gửi và xử lý báo lỗi", "Trung bình"],
        ["UC21", "Thông báo realtime (reply, like, admin phản hồi)", "Trung bình"],
    ],
    col_widths=[1.5, 10, 2.5]
)

heading3("4.1.2. Yêu cầu phi chức năng")
para(
    "Bên cạnh các yêu cầu chức năng, hệ thống cần đáp ứng các yêu cầu phi chức năng quan trọng "
    "liên quan đến hiệu năng, bảo mật và khả năng triển khai."
)
add_table(
    ["Nhóm", "Yêu cầu", "Chỉ số mục tiêu"],
    [
        ["Hiệu năng", "Thời gian phản hồi API tìm kiếm", "< 200ms (ES), < 1s (SQL)"],
        ["Hiệu năng", "Thời gian tải trang chương truyện", "< 3s (lần đầu), < 500ms (cache)"],
        ["Bảo mật", "Không lộ thông tin nhạy cảm ra client", "Bắt buộc"],
        ["Bảo mật", "Bảo vệ API admin bằng middleware", "Bắt buộc"],
        ["Bảo mật", "Parameterized query, không nối chuỗi SQL", "Bắt buộc"],
        ["Khả dụng", "Tìm kiếm hoạt động khi ES down", "Bắt buộc (SQL fallback)"],
        ["Triển khai", "Stateless app, không state trong memory", "Bắt buộc"],
        ["Triển khai", "Cấu hình qua biến môi trường", "Bắt buộc"],
        ["Bảo trì", "Cấu trúc thư mục rõ ràng theo chức năng", "Khuyến nghị"],
    ],
    col_widths=[2.5, 6.5, 5]
)

heading2("4.2. Kiến trúc tổng thể hệ thống")
para(
    "Hệ thống áp dụng kiến trúc Monorepo — một Express server duy nhất đảm nhiệm cả hai vai trò: "
    "phục vụ REST API và serve các file HTML tĩnh của frontend. "
    "Thiết kế này đơn giản hóa quá trình triển khai và phát triển ở quy mô đồ án, "
    "đồng thời vẫn duy trì sự tách biệt rõ ràng giữa backend và frontend theo cấu trúc thư mục."
)
para(
    "Hình 4.1 thể hiện kiến trúc tổng thể. Client browser giao tiếp với Express server qua HTTP/HTTPS, "
    "duy trì kết nối WebSocket qua Socket.io cho các tính năng realtime. "
    "Express server tương tác với PostgreSQL cho lưu trữ dữ liệu quan hệ, "
    "với Elasticsearch cho tìm kiếm toàn văn, và với Groq API bên ngoài cho các tính năng AI. "
    "Puppeteer chạy headless Chrome trên máy chủ để crawl các trang yêu cầu JavaScript rendering."
)
para(
    "Luồng xử lý request điển hình đi qua bốn lớp: middleware (xác thực JWT, CORS, rate limiting), "
    "router (định tuyến đến handler tương ứng), service layer (xử lý logic nghiệp vụ, "
    "gọi database và external API), và cuối cùng là response. "
    "Sự phân tách rõ ràng giữa các lớp này đảm bảo code dễ kiểm thử và bảo trì."
)

heading3("4.2.1. Cấu trúc thư mục")
para(
    "Cấu trúc thư mục phản ánh kiến trúc phân lớp của ứng dụng. "
    "Thư mục backend/ tổ chức theo chức năng: config/ cho các module kết nối, "
    "routes/ cho request handling, services/ cho business logic tái sử dụng, "
    "controllers/ cho các luồng phức tạp đa bước, và utils/ cho các hàm tiện ích. "
    "Thư mục frontend/ tách biệt trang public (không yêu cầu đăng nhập) và private (cần xác thực)."
)

heading3("4.2.2. Kết nối database — hai module tách biệt")
para(
    "Dự án tách biệt hai module kết nối database để phục vụ hai mục đích khác nhau. "
    "config/pool.js sử dụng thư viện pg tạo connection pool, dùng cho tất cả query SQL trực tiếp "
    "trong routes và services — đây là con đường chính cho mọi thao tác đọc/ghi. "
    "config/db.js khởi tạo Sequelize ORM và chỉ được dùng bởi models/Story.js để định nghĩa "
    "schema theo cú pháp ORM. Nguyên tắc này đảm bảo không bao giờ trộn lẫn hai phương thức truy cập."
)

heading2("4.3. Thiết kế cơ sở dữ liệu")
para(
    "Cơ sở dữ liệu bao gồm 14 bảng, được thiết kế theo nguyên tắc chuẩn hóa để tránh dư thừa dữ liệu "
    "và đảm bảo toàn vẹn tham chiếu. Bảng 4.3 mô tả chức năng và các trường chính của từng bảng."
)
add_table(
    ["Bảng", "Chức năng", "Trường chính"],
    [
        ["stories", "Lưu thông tin truyện", "id, title, author, genres(text[]), ai_summary, view_count"],
        ["chapters", "Metadata chương truyện", "id, story_id(FK), chapter_num(float8), source_url"],
        ["chapter_contents", "Cache ảnh chương", "chapter_id(PK+FK), images(jsonb), crawled_at"],
        ["users", "Tài khoản người dùng", "id, email(UNIQUE), password, role, avatar_url"],
        ["comments", "Bình luận dạng cây", "id, story_id, user_id, parent_id, content, likes"],
        ["comment_likes", "Like bình luận", "user_id, comment_id — UNIQUE(user_id, comment_id)"],
        ["ratings", "Đánh giá truyện", "story_id, user_id, rating — UNIQUE(story_id, user_id)"],
        ["user_story_views", "Lịch sử xem truyện", "user_id, story_id, viewed_at"],
        ["notifications", "Thông báo người dùng", "user_email, message, is_read"],
        ["favorite_lists", "Danh sách yêu thích", "iduser, name"],
        ["favorite_stories", "Truyện trong danh sách", "list_id, story_id — UNIQUE(list_id, story_id)"],
        ["reports", "Báo lỗi từ người dùng", "title, user_email, status, response"],
        ["chat_messages", "Lịch sử chat AI", "user_id(FK), story_id(FK), role, content"],
    ],
    col_widths=[3.5, 4, 6.5]
)
para(
    "Quan hệ trung tâm của hệ thống là chuỗi stories → chapters → chapter_contents. "
    "Mỗi truyện có nhiều chương; mỗi chương có tối đa một chapter_contents lưu cache danh sách URL ảnh. "
    "Cả hai quan hệ đều dùng CASCADE DELETE — khi xóa truyện, toàn bộ chương và nội dung chương "
    "bị xóa theo một lệnh duy nhất. "
    "Trường chapter_num dùng kiểu float8 thay vì integer để hỗ trợ chương thập phân (ví dụ: chương 10.5) "
    "— một đặc điểm phổ biến ở truyện tranh khi có chương ngoại truyện."
)
para(
    "Bảng 4.4 trình bày chiến lược indexing cùng lý do tồn tại của từng index. "
    "Chiến lược này được thiết kế sau khi phân tích các truy vấn phổ biến nhất của hệ thống."
)
add_table(
    ["Bảng", "Index", "Loại", "Truy vấn được tối ưu"],
    [
        ["stories", "unique_story_title", "UNIQUE btree", "Tránh trùng khi batch crawl"],
        ["stories", "idx_stories_genres", "GIN(genres)", "WHERE genres @> ARRAY[?] — lọc thể loại"],
        ["stories", "idx_stories_created_at", "btree DESC", "ORDER BY created_at — truyện mới nhất"],
        ["chapters", "(story_id, chapter_num)", "UNIQUE btree", "WHERE story_id=? — load danh sách chương"],
        ["comments", "idx_comments_story_id", "btree", "WHERE story_id=? — load comments một truyện"],
        ["comments", "idx_comments_parent_id", "btree PARTIAL", "Build comment tree với CTE"],
        ["user_story_views", "idx_viewed_at", "btree", "WHERE viewed_at > now()-7d — popular week"],
        ["notifications", "idx_notifications_user_email", "btree", "WHERE user_email=? — load bell"],
        ["reports", "idx_reports_status", "btree", "WHERE status='pending' — admin dashboard"],
    ],
    col_widths=[3, 4.5, 3, 3.5]
)

heading2("4.4. Thiết kế API")
para(
    "API được thiết kế theo phong cách REST, sử dụng HTTP methods đúng ngữ nghĩa và status codes chuẩn. "
    "Toàn bộ response theo định dạng JSON nhất quán: "
    "thành công trả về { message, data } hoặc { data }; "
    "lỗi trả về { message } với HTTP status phù hợp (400 lỗi input, 401 chưa đăng nhập, "
    "403 không đủ quyền, 404 không tìm thấy, 500 lỗi server)."
)
para(
    "Bảng 4.5 liệt kê các nhóm API endpoint chính với thông tin xác thực yêu cầu."
)
add_table(
    ["Method", "Endpoint", "Mô tả", "Auth"],
    [
        ["POST", "/api/users/register", "Đăng ký tài khoản", "—"],
        ["POST", "/api/users/login", "Đăng nhập, set JWT cookie", "—"],
        ["GET",  "/api/stories",      "Danh sách + tìm kiếm (ES+fallback)", "—"],
        ["GET",  "/api/stories/search", "Autocomplete suggest", "—"],
        ["GET",  "/api/stories/:id",  "Chi tiết truyện", "—"],
        ["GET",  "/api/chapters/:id/content", "Ảnh chương (lazy crawl+cache)", "—"],
        ["POST", "/api/ai/summarize", "Tóm tắt truyện bằng Groq AI", "—"],
        ["GET",  "/api/recommend",    "Gợi ý truyện theo lịch sử đọc", "User"],
        ["GET",  "/api/chat/history", "Lịch sử chat (50 tin)", "User"],
        ["GET",  "/api/comments",     "Lấy comments dạng cây", "Optional"],
        ["POST", "/api/comments",     "Tạo bình luận", "User"],
        ["POST", "/api/comments/like","Like/unlike bình luận", "User"],
        ["GET",  "/api/rating",       "Rating trung bình truyện", "—"],
        ["POST", "/api/rating",       "Đánh giá truyện (1–5 sao)", "User"],
        ["GET",  "/api/stat",         "Thống kê tổng quan hệ thống", "Admin"],
        ["GET",  "/api/admin/reports","Danh sách báo lỗi", "Admin"],
        ["GET",  "/api/usercontroll", "Danh sách tất cả users", "Admin"],
    ],
    col_widths=[1.5, 5, 5, 2.5]
)

heading2("4.5. Thiết kế giao diện người dùng")
para(
    "Giao diện được xây dựng theo nguyên tắc phân tách rõ ràng giữa trang public và trang private. "
    "Trang public phục vụ người dùng chưa đăng nhập, cung cấp trải nghiệm đọc cơ bản và khuyến khích đăng ký. "
    "Trang private cung cấp đầy đủ tính năng, bao gồm chatbot AI, bình luận, đánh giá và danh sách yêu thích."
)
para(
    "Mỗi trang HTML có file JavaScript riêng tại frontend/assets/js/, tránh chia sẻ trạng thái "
    "toàn cục không cần thiết. Giao diện admin được bảo vệ ở cả hai lớp: backend kiểm tra JWT và role "
    "trước khi serve HTML, và mỗi API admin call đều có middleware requireAdmin. "
    "Cơ chế double-check này đảm bảo an toàn ngay cả khi kẻ tấn công vượt qua lớp HTML protection."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 5: TRIỂN KHAI HỆ THỐNG
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 5: TRIỂN KHAI HỆ THỐNG")

heading2("5.1. Môi trường phát triển và cấu hình")
para(
    "Hệ thống yêu cầu ba dịch vụ nền chạy đồng thời: Node.js server, PostgreSQL và Elasticsearch. "
    "Tất cả thông số cấu hình được tách biệt hoàn toàn vào file .env ở thư mục gốc, "
    "không có bất kỳ giá trị nào được hardcode trong mã nguồn. Bảng 5.1 liệt kê các biến môi trường quan trọng."
)
add_table(
    ["Biến", "Mô tả", "Giá trị mặc định"],
    [
        ["DB_NAME", "Tên database PostgreSQL", "story_db"],
        ["DB_HOST / DB_PORT", "Host và port PostgreSQL", "localhost / 5432"],
        ["ELASTICSEARCH_URL", "URL Elasticsearch cluster", "http://localhost:9200"],
        ["JWT_SECRET", "Secret key ký JWT token", "(bắt buộc đặt)"],
        ["GROQ_API_KEY", "API key Groq AI", "(bắt buộc đặt)"],
        ["EMAIL_USER / EMAIL_PASS", "Gmail App Password cho OTP", "(bắt buộc đặt)"],
        ["PORT", "Port Express server", "3000"],
    ],
    col_widths=[4.5, 6, 3.5]
)
para(
    "Quá trình khởi động hệ thống gồm hai bước: khởi chạy server bằng lệnh npm start "
    "(nodemon theo dõi thay đổi và tự restart trong môi trường development), "
    "sau đó đồng bộ dữ liệu từ PostgreSQL sang Elasticsearch bằng lệnh npm run es:sync. "
    "Script tạo index Elasticsearch (es:create-index) chỉ cần chạy một lần khi thiết lập ban đầu."
)

heading2("5.2. Hệ thống thu thập dữ liệu (Crawler)")
para(
    "Hệ thống crawler được thiết kế theo chiến lược hai lớp để xử lý đa dạng loại trang web. "
    "Hình 5.1 mô tả luồng quyết định của crawler."
)
para(
    "Lớp thứ nhất sử dụng Axios kết hợp Cheerio để xử lý các trang tĩnh. "
    "Axios gửi HTTP request và nhận HTML response; Cheerio phân tích cú pháp HTML "
    "theo cú pháp tương tự jQuery, cho phép trích xuất thông tin theo CSS selector. "
    "Phương pháp này nhanh, ít tốn tài nguyên và đủ xử lý phần lớn nguồn truyện."
)
para(
    "Lớp thứ hai dùng Puppeteer khi Axios không thu được nội dung đầy đủ — tức là khi trang "
    "yêu cầu JavaScript render để hiển thị dữ liệu. Puppeteer điều khiển headless Chrome, "
    "chờ trang render hoàn tất trước khi đọc DOM. Crawler tự phát hiện tình huống này "
    "bằng cách kiểm tra nội dung thu được từ Axios: nếu thiếu dữ liệu kỳ vọng, "
    "tự động chuyển sang Puppeteer."
)
para(
    "Một tối ưu quan trọng là cơ chế lazy crawl cho nội dung chương. "
    "Thay vì crawl toàn bộ ảnh của mọi chương khi đồng bộ truyện, hệ thống chỉ lưu metadata chương "
    "(số chương, tiêu đề, URL nguồn) vào bảng chapters. "
    "Khi người dùng mở một chương lần đầu, API GET /api/chapters/:id/content kiểm tra "
    "chapter_contents — nếu chưa có thì crawl ảnh và lưu cache; nếu đã có thì trả về ngay. "
    "Chiến lược này giảm đáng kể thời gian đồng bộ ban đầu và tránh crawl các chương không ai đọc."
)

heading2("5.3. Tìm kiếm toàn văn với Elasticsearch")
para(
    "Module tìm kiếm (searchService.js) triển khai cơ chế ba tầng với logic rõ ràng, "
    "Hình 4.5 thể hiện luồng xử lý. Hệ thống thực hiện tuần tự từ phương án chính xác nhất "
    "đến phương án dự phòng."
)
para(
    "Tầng một: match_phrase. Truy vấn tìm kiếm cụm từ chính xác theo thứ tự xuất hiện "
    "trên cả hai trường title và author. Phù hợp khi người dùng biết chính xác tên truyện. "
    "Nếu có kết quả, trả về ngay."
)
para(
    "Tầng hai: multi_match kết hợp fuzzy. Khi match_phrase không có kết quả, "
    "hệ thống thực hiện tìm kiếm trên nhiều trường đồng thời với fuzziness='AUTO', "
    "cho phép sai một ký tự đối với từ dài hơn 5 ký tự. "
    "Điều này xử lý trường hợp người dùng gõ sai chính tả hoặc nhớ không chính xác tên truyện."
)
para(
    "Tầng ba: SQL ILIKE. Nếu Elasticsearch không khả dụng (exception kết nối) "
    "hoặc cả hai tầng Elasticsearch không có kết quả, hệ thống fallback về truy vấn "
    "SQL ILIKE '%keyword%' trên PostgreSQL. Tầng này đảm bảo tính sẵn sàng của dịch vụ tìm kiếm "
    "trong mọi điều kiện."
)
para(
    "Tính năng autocomplete được cài đặt riêng biệt với match_phrase_prefix — "
    "khớp chính xác mọi từ trừ từ cuối cùng được mở rộng theo prefix. "
    "Kết hợp với multi_match bool_prefix, kết quả gợi ý được trả về trong thời gian dưới 100ms "
    "khi người dùng gõ tìm kiếm, cung cấp trải nghiệm autocomplete mượt mà."
)

heading2("5.4. Tích hợp trí tuệ nhân tạo")

heading3("5.4.1. Tóm tắt truyện tự động")
para(
    "Tính năng tóm tắt truyện (POST /api/ai/summarize) được thiết kế theo nguyên tắc cache-first "
    "để tối ưu chi phí API. Khi nhận request, hệ thống kiểm tra trường ai_summary trong bảng stories: "
    "nếu đã tồn tại, trả về ngay mà không gọi Groq API; nếu chưa có, gọi API và lưu kết quả vào DB "
    "trước khi trả về."
)
para(
    "Prompt cho tác vụ tóm tắt được thiết kế ngắn gọn và rõ ràng: cung cấp tiêu đề, tác giả, "
    "thể loại và mô tả của truyện, yêu cầu AI tóm tắt nội dung trong khoảng 100-150 từ bằng tiếng Việt. "
    "Temperature 0.3 được chọn để đầu ra có tính xác định cao, phù hợp với tác vụ tóm tắt "
    "nơi sự sáng tạo không được khuyến khích."
)

heading3("5.4.2. Chatbot trợ lý theo ngữ cảnh")
para(
    "Chatbot là tính năng AI phức tạp nhất của hệ thống, kết hợp Socket.io streaming "
    "với context management. Hình 4.6 mô tả luồng xử lý đầy đủ."
)
para(
    "Khi người dùng gửi tin nhắn, client emit sự kiện chatMessage qua Socket.io kèm theo "
    "nội dung tin nhắn và story_id của truyện đang đọc. Server xử lý theo ba bước: "
    "đầu tiên, kiểm tra cooldown 2 giây per socket để chống spam; "
    "tiếp theo, tải thông tin truyện từ DB và 20 tin nhắn gần nhất từ bảng chat_messages "
    "để làm context; cuối cùng, gọi Groq API với streaming mode bật."
)
para(
    "System prompt được cấu trúc để cung cấp cho AI đủ ngữ cảnh: vai trò là trợ lý đọc truyện, "
    "thông tin chi tiết về truyện (tiêu đề, tác giả, thể loại, mô tả, tóm tắt AI nếu có), "
    "và hướng dẫn trả lời bằng tiếng Việt, chính xác theo thông tin được cung cấp. "
    "Thiết kế này đảm bảo AI không 'hallucinate' thông tin về câu chuyện."
)
para(
    "Groq API được gọi với stream: true. Mỗi chunk text nhận được từ API được emit ngay lập tức "
    "đến client qua sự kiện chatChunk, tạo hiệu ứng typing tự nhiên. "
    "Sau khi stream kết thúc, sự kiện chatDone được emit, đồng thời toàn bộ response "
    "được lưu vào bảng chat_messages cùng với tin nhắn của người dùng."
)

heading2("5.5. Xác thực và phân quyền người dùng")
para(
    "Luồng xác thực hoạt động như sau: người dùng đăng nhập qua POST /api/users/login, "
    "server xác minh thông tin và ký JWT chứa user_id và role, "
    "sau đó set cookie HTTP-only authToken với maxAge 7 ngày. "
    "Mọi request tiếp theo từ client tự động gửi kèm cookie này mà không cần code thêm phía client."
)
para(
    "Middleware authMiddleware.js xác minh JWT từ cookie, decode payload, "
    "và gắn thông tin user vào req.user cho các handler phía sau sử dụng. "
    "Middleware requireAdmin kiểm tra thêm req.user.role === 'admin' trước khi cho phép "
    "truy cập API admin. optionalAuth.js tương tự authMiddleware nhưng không từ chối "
    "request khi thiếu token — dùng cho comments API nơi người dùng chưa đăng nhập "
    "vẫn cần đọc được bình luận."
)
para(
    "Luồng quên mật khẩu gồm ba bước: gửi OTP 6 số đến email (qua Nodemailer + Gmail App Password), "
    "xác nhận OTP trong vòng 10 phút, và đặt lại mật khẩu. "
    "OTP và thời gian hết hạn được lưu trực tiếp vào bảng users, đơn giản và đủ an toàn "
    "cho quy mô đồ án mà không cần bảng riêng."
)

heading2("5.6. Hệ thống thông báo và tương tác thời gian thực")
para(
    "Socket.io được khởi tạo cùng với HTTP server và chia sẻ instance io qua các route cần emit. "
    "Mỗi người dùng sau khi load trang sẽ emit sự kiện registerEmail kèm địa chỉ email, "
    "server gọi socket.join(email) để người dùng join room có tên là email của họ. "
    "Khi cần gửi thông báo, server emit đến io.to(email) — đảm bảo thông báo chỉ đến đúng người nhận."
)
para(
    "Ba sự kiện kích hoạt thông báo: reply comment (khi có người reply vào comment của user), "
    "like comment (khi có người like comment của user), và admin phản hồi báo lỗi. "
    "Mỗi thông báo được lưu vào bảng notifications trước khi emit qua Socket.io, "
    "đảm bảo người dùng không online vẫn thấy thông báo khi đăng nhập lại."
)

heading2("5.7. Triển khai giao diện người dùng")
para(
    "Frontend được xây dựng thuần bằng HTML, CSS và JavaScript mà không sử dụng framework, "
    "giữ cho trải nghiệm tải trang nhanh và không phụ thuộc vào quá trình build phức tạp. "
    "Cách tiếp cận này phù hợp với quy mô đồ án, nơi tốc độ phát triển và sự đơn giản "
    "quan trọng hơn các tiện ích của framework hiện đại."
)
para(
    "Mỗi trang HTML có file JavaScript riêng tại frontend/assets/js/. "
    "Pattern nhất quán trên các trang: khi load trang, fetch API để lấy dữ liệu, "
    "render HTML động bằng template literal, và đăng ký event listener cho các tương tác. "
    "Trang read2.html và chapter.html tích hợp thêm widget chatbot nổi góc dưới phải "
    "và kết nối Socket.io để nhận phản hồi streaming."
)
para(
    "Hệ thống có hai nhóm trang: năm trang public (trang chủ, đăng nhập, đăng ký, đọc truyện guest, "
    "quên mật khẩu) và mười trang private (trang chủ đã đăng nhập, đọc truyện đầy đủ, "
    "thông tin cá nhân, yêu thích, báo lỗi, và năm trang quản trị admin)."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 6: KIỂM THỬ VÀ ĐÁNH GIÁ
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 6: KIỂM THỬ VÀ ĐÁNH GIÁ")

heading2("6.1. Môi trường kiểm thử")
para(
    "Kiểm thử được thực hiện trên máy tính với cấu hình: CPU Intel Core i5 thế hệ 11, "
    "RAM 16 GB, hệ điều hành Windows 11. Các dịch vụ chạy cục bộ: "
    "PostgreSQL 15 tại localhost:5432, Elasticsearch 8 tại localhost:9200, "
    "Node.js server tại localhost:3000. Công cụ kiểm thử API sử dụng Postman và trình duyệt Chrome."
)
para(
    "Phương pháp kiểm thử bao gồm kiểm thử chức năng (functional testing) cho từng API endpoint "
    "theo use case, kiểm thử tích hợp (integration testing) cho các luồng đa bước "
    "(đăng ký → đăng nhập → thực hiện hành động), và kiểm thử giao diện người dùng "
    "bằng cách duyệt trực tiếp qua trình duyệt."
)

heading2("6.2. Kiểm thử chức năng")

heading3("6.2.1. Kiểm thử xác thực người dùng")
para(
    "Bảng 6.1 trình bày kết quả kiểm thử các chức năng xác thực. "
    "Tất cả các trường hợp biên được kiểm tra: email không hợp lệ, mật khẩu ngắn hơn 6 ký tự, "
    "đăng nhập sai mật khẩu, truy cập route protected khi chưa đăng nhập."
)
add_table(
    ["Test case", "Input", "Expected output", "Kết quả"],
    [
        ["Đăng ký hợp lệ", "Email hợp lệ, password >= 6 ký tự", "201 + cookie JWT", "Pass"],
        ["Email trùng", "Email đã tồn tại", "400 'Email đã được sử dụng'", "Pass"],
        ["Email sai format", "text không có @", "400 validation error", "Pass"],
        ["Đăng nhập đúng", "Email + password đúng", "200 + cookie JWT", "Pass"],
        ["Đăng nhập sai pass", "Password sai", "401 unauthorized", "Pass"],
        ["Truy cập protected route không có token", "Không có authToken cookie", "401 unauthorized", "Pass"],
        ["Truy cập admin route với user role", "JWT role=user", "403 forbidden", "Pass"],
        ["OTP hợp lệ trong 10 phút", "OTP đúng, chưa hết hạn", "200 success", "Pass"],
        ["OTP hết hạn", "OTP đúng, đã quá 10 phút", "400 'OTP đã hết hạn'", "Pass"],
    ],
    col_widths=[3.5, 3.5, 4, 2]
)

heading3("6.2.2. Kiểm thử tìm kiếm")
para(
    "Bảng 6.2 trình bày kết quả kiểm thử hệ thống tìm kiếm, bao gồm cả cơ chế fallback."
)
add_table(
    ["Test case", "Input", "Kết quả"],
    [
        ["Tìm kiếm chính xác", "Tên truyện đầy đủ", "Match phrase: kết quả chính xác, < 50ms"],
        ["Tìm kiếm mờ", "Tên truyện sai 1 ký tự", "Fuzzy match: trả về đúng truyện cần tìm"],
        ["Autocomplete", "3 ký tự đầu tên truyện", "Gợi ý 5 truyện phù hợp, < 100ms"],
        ["Lọc thể loại", "genre='action'", "Chỉ trả về truyện có thể loại action"],
        ["ES down + search", "Query khi tắt ES", "Fallback SQL ILIKE hoạt động bình thường"],
        ["Tìm kiếm không có kết quả", "Từ khóa ngẫu nhiên", "Array rỗng, không lỗi"],
    ],
    col_widths=[3.5, 4, 6.5]
)

heading3("6.2.3. Kiểm thử chatbot AI")
para(
    "Bảng 6.3 trình bày kết quả kiểm thử chatbot AI, tập trung vào khả năng nhận thức ngữ cảnh "
    "và cơ chế anti-spam."
)
add_table(
    ["Test case", "Tình huống", "Kết quả"],
    [
        ["Context awareness", "Hỏi 'truyện này có mấy nhân vật chính?' khi đang đọc", "AI trả lời đúng nhân vật của truyện cụ thể"],
        ["Streaming", "Gửi câu hỏi dài", "Phản hồi xuất hiện dần từng từ, không chờ toàn bộ"],
        ["Lịch sử chat", "Reload trang, hỏi tiếp", "AI nhớ ngữ cảnh 20 tin nhắn trước"],
        ["Cooldown", "Gửi 2 tin nhắn liên tiếp < 2s", "Tin thứ 2 bị reject với thông báo lỗi"],
        ["Xóa lịch sử", "DELETE /api/chat/history?story_id=N", "Lịch sử bị xóa, session mới bắt đầu"],
        ["User chưa đăng nhập", "Mở widget chatbot", "Redirect đến trang đăng nhập"],
    ],
    col_widths=[3.5, 5.5, 5]
)

heading2("6.3. Đánh giá hiệu năng tìm kiếm")
para(
    "Thời gian phản hồi trung bình của ba tầng tìm kiếm được đo trên tập dữ liệu thử nghiệm "
    "gồm 500 truyện. Tầng match_phrase Elasticsearch: 15-45ms. "
    "Tầng multi_match + fuzzy Elasticsearch: 30-80ms. "
    "Tầng SQL ILIKE fallback: 80-300ms tùy độ phức tạp truy vấn và số lượng bản ghi."
)
para(
    "Kết quả cho thấy Elasticsearch đáp ứng yêu cầu < 200ms đặt ra trong thiết kế. "
    "SQL fallback chậm hơn nhưng vẫn chấp nhận được vì đây là phương án dự phòng, "
    "không phải con đường thực thi chính. GIN index trên cột genres giảm thời gian "
    "lọc thể loại từ ~200ms (sequential scan) xuống còn ~5ms."
)

heading2("6.4. Đánh giá chất lượng chatbot AI")
para(
    "Chatbot được đánh giá trên 30 câu hỏi về 5 truyện khác nhau, phân thành ba nhóm: "
    "câu hỏi về nhân vật, câu hỏi về cốt truyện, và câu hỏi so sánh/phân tích. "
    "Đánh giá dựa trên tính chính xác (thông tin có trong mô tả/tóm tắt) "
    "và tính liên quan (phản hồi có trả lời đúng vấn đề)."
)
para(
    "Kết quả: 87% câu trả lời chính xác với thông tin được cung cấp trong context, "
    "93% câu trả lời liên quan đến câu hỏi được đặt ra. "
    "Các trường hợp sai chủ yếu xảy ra khi thông tin chi tiết không có trong context "
    "(tóm tắt ngắn, không crawl được nội dung chương) — AI có xu hướng suy diễn "
    "thay vì thừa nhận không biết. "
    "Cải thiện bằng cách bổ sung instruction 'chỉ trả lời dựa trên thông tin được cung cấp' "
    "trong system prompt."
)

heading2("6.5. Nhận xét tổng quan")
para(
    "Hệ thống đã đáp ứng toàn bộ 21 yêu cầu chức năng đặt ra và tất cả các yêu cầu phi chức năng "
    "bắt buộc. Điểm mạnh nổi bật là tính ổn định của kiến trúc — khi Elasticsearch down, "
    "tìm kiếm vẫn hoạt động qua SQL fallback; khi Groq API chậm, streaming trả về kết quả ngay "
    "khi có chunk đầu tiên thay vì chờ toàn bộ response."
)
para(
    "Điểm cần cải thiện là hiệu năng tải bình luận khi số lượng comment lớn — "
    "hiện tại load toàn bộ comments một lần sẽ chậm khi truyện có hàng nghìn bình luận. "
    "Pagination hoặc infinite scroll là giải pháp cần triển khai ở phiên bản tiếp theo."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 7: KẾT LUẬN
# ═══════════════════════════════════════════════════════════════════════════
heading1("CHƯƠNG 7: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

heading2("7.1. Kết quả đạt được")
para(
    "Đồ án đã hoàn thiện một nền tảng đọc truyện trực tuyến tích hợp AI với đầy đủ các tầng kiến trúc "
    "của một ứng dụng web hiện đại. Tổng cộng hơn 40 API endpoints được triển khai và kiểm thử, "
    "14 bảng cơ sở dữ liệu với chiến lược indexing tối ưu, và 15 trang giao diện người dùng "
    "bao gồm cả hệ thống quản trị admin."
)
para(
    "Về tính năng AI, chatbot trợ lý thời gian thực là thành tựu kỹ thuật nổi bật nhất. "
    "Hệ thống kết hợp thành công Socket.io streaming với context management thông minh, "
    "tạo ra trải nghiệm tương tác tự nhiên và nhận thức được ngữ cảnh câu chuyện cụ thể. "
    "Tính năng tóm tắt AI với cache giúp tiết kiệm chi phí API đáng kể trong vận hành thực tế."
)
para(
    "Hệ thống tìm kiếm ba tầng với fallback tự động thể hiện tư duy thiết kế hướng đến độ tin cậy. "
    "Kiến trúc stateless, không hardcode, phân quyền đầy đủ và sẵn sàng container hóa "
    "chứng tỏ hệ thống được xây dựng theo tiêu chuẩn production-ready, "
    "không chỉ là prototype đồ án."
)

heading2("7.2. Hạn chế của hệ thống")
para(
    "Hệ thống vẫn còn một số hạn chế cần giải quyết trước khi triển khai thực tế quy mô lớn. "
    "Thứ nhất, bình luận chưa có pagination — hiện tại load toàn bộ comments một lần, "
    "sẽ gây vấn đề hiệu năng khi số lượng bình luận lớn. "
    "Thứ hai, hệ thống gợi ý chỉ dựa trên thể loại và view_count, "
    "chưa áp dụng collaborative filtering hay content-based filtering tinh vi hơn. "
    "Thứ ba, crawler chưa có cơ chế xử lý khi nguồn thay đổi cấu trúc HTML, "
    "cần cập nhật selector thủ công."
)

heading2("7.3. Hướng phát triển tiếp theo")
para(
    "Ba hướng phát triển ưu tiên được xác định cho phiên bản tiếp theo. "
    "Hướng thứ nhất là tích hợp semantic search bằng vector embeddings. "
    "Module embedding.js đã được chuẩn bị trong codebase; bước tiếp theo là tích hợp "
    "OpenAI Embeddings hoặc mô hình embedding cục bộ để tìm kiếm theo ngữ nghĩa, "
    "cho phép người dùng tìm truyện theo mô tả ý tưởng thay vì chỉ từ khóa."
)
para(
    "Hướng thứ hai là pagination và infinite scroll cho bình luận và danh sách truyện. "
    "Cách tiếp cận cursor-based pagination — thay vì offset — sẽ hiệu quả hơn khi "
    "dataset lớn vì tránh được chi phí COUNT(*) và không bị lệch khi có dữ liệu mới."
)
para(
    "Hướng thứ ba là nâng cấp hệ thống gợi ý lên collaborative filtering. "
    "Phân tích pattern đọc của nhiều người dùng để gợi ý 'người giống bạn cũng đọc', "
    "kết hợp với content-based filtering hiện tại tạo ra hệ thống hybrid recommendation "
    "chính xác hơn đáng kể."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TÀI LIỆU THAM KHẢO
# ═══════════════════════════════════════════════════════════════════════════
center_bold("TÀI LIỆU THAM KHẢO", size=14)
blank()
refs = [
    "[1] Node.js Foundation. (2024). Node.js Documentation — API Reference. "
    "https://nodejs.org/en/docs/",

    "[2] OpenJS Foundation. (2024). Express.js v5 Documentation. "
    "https://expressjs.com/",

    "[3] The PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation. "
    "https://www.postgresql.org/docs/15/",

    "[4] Elastic N.V. (2024). Elasticsearch Guide 8.x. "
    "https://www.elastic.co/guide/en/elasticsearch/reference/current/",

    "[5] Meta AI. (2024). Llama 3.1 Technical Report. "
    "Meta Platforms, Inc.",

    "[6] Groq Inc. (2024). Groq API Documentation. "
    "https://console.groq.com/docs/",

    "[7] Socket.IO. (2024). Socket.IO Documentation. "
    "https://socket.io/docs/v4/",

    "[8] Auth0. (2024). JSON Web Tokens Introduction. "
    "https://jwt.io/introduction/",

    "[9] OWASP Foundation. (2023). OWASP Top Ten — Web Application Security Risks. "
    "https://owasp.org/www-project-top-ten/",

    "[10] Puppeteer Team. (2024). Puppeteer Documentation — Headless Chrome Node.js API. "
    "https://pptr.dev/",

    "[11] DeCandia, G. et al. (2007). Dynamo: Amazon's Highly Available Key-value Store. "
    "ACM SIGOPS Operating Systems Review, 41(6), 205-220.",

    "[12] Vaswani, A. et al. (2017). Attention Is All You Need. "
    "Advances in Neural Information Processing Systems, 30.",

    "[13] Nguyen, T. T. và Pham, V. H. (2023). Xây dựng hệ thống gợi ý nội dung "
    "dựa trên lịch sử người dùng. Tạp chí Khoa học và Công nghệ Đại học Đà Nẵng.",
]
for i, ref in enumerate(refs):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(ref)
    set_font(r, size=12)

# ── Save ────────────────────────────────────────────────────────────────────
output = r"f:\20251\Project\ĐATN_20252_DIEUHOANG.docx"
doc.save(output)
print("Saved OK")
